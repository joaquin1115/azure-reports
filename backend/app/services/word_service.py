import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.services.analisis_service import ResultadoMetrica

COLOR_SUBTITULO = RGBColor(0x19, 0x87, 0xAF)


def _set_font(run, size_pt: int, color: RGBColor = RGBColor(0, 0, 0)):
    run.font.name = "Segoe UI Semilight"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI Semilight")
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    run = p.add_run(text)
    size = 16 if level == 1 else 14 if level == 2 else 12
    _set_font(run, size, COLOR_SUBTITULO)
    run.bold = True


def _add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, 10)


def _asignar_bordes_tabla(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '4')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), 'D9D9D9')
        borders.append(elem)
    tblPr.append(borders)

def _detalle_observaciones(metricas: list[tuple[str, ResultadoMetrica | None]]) -> str:
    detalles = []
    for titulo, metrica in metricas:
        if not metrica:
            continue
        observaciones = [obs.strip() for obs in (metrica.observaciones or []) if obs and obs.strip()]
        if observaciones:
            detalles.append(f"{titulo}: {' '.join(observaciones)}")
    return " | ".join(detalles) if detalles else "Sin observaciones"


def _buscar_metrica(metricas: list[ResultadoMetrica], aliases: list[str]) -> ResultadoMetrica | None:
    if not metricas:
        return None
    normalized = [a.lower() for a in aliases]
    for m in metricas:
        nombre = m.nombre.lower()
        if any(a in nombre for a in normalized):
            return m
    return None


def _tipo_recurso(valor) -> str:
    raw = getattr(valor, "value", valor)
    return str(raw).strip().upper()


def _fmt_rango(m: ResultadoMetrica | None) -> str:
    if not m or not m.valores:
        return "Sin datos"
    return f"{m.minimo:.2f}% - {m.maximo:.2f}%"


def _label_tipo(tipo: str) -> str:
    return {
        "VM": "Máquina Virtual",
        "DB": "Base de Datos SQL",
        "ASP": "App Service Plan",
    }.get(tipo, tipo or "Recurso")


def _metricas_por_tipo(tipo: str, metricas: list[ResultadoMetrica]) -> list[tuple[str, ResultadoMetrica | None]]:
    if tipo == "VM":
        return [
            ("Porcentaje de Uso de CPU", _buscar_metrica(metricas, ["percentage cpu", "cpupercentage", "cpu"])),
            ("Porcentaje de Uso de Memoria RAM", _buscar_metrica(metricas, ["memorypercentage", "available memory", "memory", "memoria"])),
        ]
    if tipo == "ASP":
        return [
            ("Porcentaje de Uso de CPU", _buscar_metrica(metricas, ["cpupercentage", "cpu percentage", "cpu"])),
            ("Porcentaje de Uso de Memoria", _buscar_metrica(metricas, ["memorypercentage", "memory percentage", "memory", "memoria"])),
        ]
    if tipo == "DB":
        return [
            ("Porcentaje de Uso de DTU", _buscar_metrica(metricas, ["dtu_consumption_percent", "dtu"]) or (metricas[0] if metricas else None)),
        ]
    return [(m.nombre, m) for m in metricas]


def _agregar_tabla_tipo(doc: Document, titulo: str, filas: list[dict], titulos_metricas: list[str]):
    _add_heading(doc, titulo, 2)
    tabla = doc.add_table(rows=1, cols=4)
    hdr = tabla.rows[0].cells
    hdr[0].text = "Recurso"
    hdr[1].text = f"{titulos_metricas[0]} (Mínimo – Máximo)" if titulos_metricas else "Métrica (Mínimo – Máximo)"
    hdr[2].text = f"{titulos_metricas[1]} (Mínimo – Máximo)" if len(titulos_metricas) > 1 else "No aplica"
    hdr[3].text = "Detalle del rendimiento"
    for fila in filas:
        row = tabla.add_row().cells
        row[0].text = fila["recurso"]
        row[1].text = fila["principal"]
        row[2].text = fila["secundaria"]
        row[3].text = fila["detalle"]
    _asignar_bordes_tabla(tabla)


def _texto_plano(valor) -> str:
    if valor is None:
        return ""
    if isinstance(valor, list):
        partes = [_texto_plano(item) for item in valor]
        return ", ".join(parte for parte in partes if parte)
    if isinstance(valor, (tuple, set)):
        partes = [_texto_plano(item) for item in valor]
        return ", ".join(parte for parte in partes if parte)
    if isinstance(valor, dict):
        partes = [
            f"{clave}: {texto}"
            for clave, item in valor.items()
            if (texto := _texto_plano(item))
        ]
        return ", ".join(partes)
    return str(valor)


def _traducir(texto: str) -> str:
    texto = _texto_plano(texto)
    if not texto:
        return ""
    reglas = {
        "Ensure Geo-replication is enabled for resilience": "Asegurar que la georreplicación esté habilitada para mejorar la resiliencia",
        "Use zone-supported App Service Plan": "Usar un App Service Plan con soporte de zonas",
        "Set minimum instance count for App Service to 2": "Configurar el número mínimo de instancias del App Service en 2",
        "Enable Health check for App Service": "Habilitar Health Check para App Service",
        "HighAvailability": "Alta disponibilidad",
        "Cost": "Costo",
        "OperationalExcellence": "Excelencia operacional",
        "Performance": "Rendimiento",
        "Security": "Seguridad",
    }
    out = texto
    for en, es in reglas.items():
        out = out.replace(en, es)
    return out


def _consolidar_recomendaciones(recomendaciones: list[dict]) -> list[dict]:
    consolidado = {}
    for r in recomendaciones:
        categoria = _texto_plano(r.get("categoria", ""))
        impacto = _texto_plano(r.get("impacto", ""))
        descripcion = _texto_plano(r.get("descripcion", ""))
        accion = _texto_plano(r.get("accion", ""))
        key = (categoria, impacto, descripcion, accion)
        entry = consolidado.setdefault(key, {
            "categoria": categoria,
            "impacto": impacto,
            "descripcion": descripcion,
            "accion": accion,
            "recursos": set(),
        })
        recurso = _texto_plano(r.get("recurso") or r.get("nombre_recurso"))
        if recurso:
            entry["recursos"].add(recurso.split("/")[-1])
    out = []
    for v in consolidado.values():
        v["recursos"] = sorted(v["recursos"])
        out.append(v)
    return out


def generar_word(cliente_nombre: str, periodo_mes: int, periodo_anio: int, usuario_nombre: str,
    recomendaciones: list[dict], resultados_por_recurso: list[dict]) -> bytes:
    doc = Document()
    _add_heading(doc, "REPORTE DE CONSUMO DE AZURE", 1)
    _add_paragraph(doc, f"Cliente: {cliente_nombre}")
    _add_paragraph(doc, f"Período: {periodo_mes:02d}/{periodo_anio}")
    _add_paragraph(doc, f"Generado: {datetime.utcnow().strftime('%d/%m/%Y %H:%M UTC')} por {usuario_nombre}")

    _add_heading(doc, "PERFORMANCE DE COMPONENTES", 1)
    filas_vm = []
    filas_asp = []
    filas_db = []
    filas_otros = []
    titulos_vm = []
    titulos_asp = []
    titulos_db = []
    titulos_otros = []

    for r in resultados_por_recurso:
        tipo = _tipo_recurso(r.get("tipo", ""))
        metricas = r.get("metricas", [])
        metricas_tipo = _metricas_por_tipo(tipo, metricas)
        principal = metricas_tipo[0][1] if metricas_tipo else None
        secundaria = metricas_tipo[1][1] if len(metricas_tipo) > 1 else None
        titulos_metricas = [titulo for titulo, _ in metricas_tipo]
        fila = {
            "recurso": r.get("nombre", ""),
            "principal": _fmt_rango(principal),
            "secundaria": _fmt_rango(secundaria) if secundaria else "No aplica",
            "detalle": _detalle_observaciones(metricas_tipo),
        }
        if tipo == "VM":
            filas_vm.append(fila)
            titulos_vm = titulos_vm or titulos_metricas
        elif tipo == "ASP":
            filas_asp.append(fila)
            titulos_asp = titulos_asp or titulos_metricas
        elif tipo == "DB":
            filas_db.append(fila)
            titulos_db = titulos_db or titulos_metricas
        else:
            filas_otros.append(fila)
            titulos_otros = titulos_otros or titulos_metricas

    if filas_vm:
        _agregar_tabla_tipo(doc, "Máquinas Virtuales", filas_vm, titulos_vm)
    if filas_asp:
        _agregar_tabla_tipo(doc, "App Service Plans", filas_asp, titulos_asp)
    if filas_db:
        _agregar_tabla_tipo(doc, "Bases de Datos SQL", filas_db, titulos_db)
    if filas_otros:
        _agregar_tabla_tipo(doc, "Otros Recursos", filas_otros, titulos_otros)

    for r in resultados_por_recurso:
        tipo = _tipo_recurso(r.get("tipo", ""))
        metricas = r.get("metricas", [])
        _add_heading(doc, f"{_label_tipo(tipo)} - {r.get('nombre','')}", 2)
        for titulo_metrica, metrica in _metricas_por_tipo(tipo, metricas):
            if not metrica:
                continue
            _add_heading(doc, titulo_metrica, 3)
            if metrica.grafico_bytes:
                doc.add_picture(io.BytesIO(metrica.grafico_bytes), width=Inches(6.2))

    _add_heading(doc, "RECOMENDACIONES Y SUGERENCIAS", 1)
    recs = _consolidar_recomendaciones(recomendaciones)
    if not recs:
        _add_paragraph(doc, "No se encontraron recomendaciones para el período analizado.")
    for r in recs:
        recursos = ", ".join(r["recursos"]) if r["recursos"] else "sin recursos identificados"
        impacto = _traducir(r["impacto"])
        categoria = _traducir(r["categoria"])
        descripcion = _traducir(r["descripcion"])
        accion = _traducir(r.get("accion", ""))
        recomendacion = f" Recomendación: {accion}." if accion else ""
        texto = f"[{impacto}] {categoria}: {descripcion}{recomendacion} Recursos: {recursos}."
        _add_paragraph(doc, texto)

    for p in doc.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()
